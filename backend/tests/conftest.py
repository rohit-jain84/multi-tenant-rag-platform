import uuid

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient

from app.config import settings
from app.db.base import Base
from app.db.session import engine

# Import every model module so SQLAlchemy registers all tables on Base.metadata
# before create_all runs. `app.models.__init__` re-exports Tenant, ApiKey,
# Document, QueryLog, and EvalResult, which is enough to ensure every mapper
# is imported. Without this import, only models transitively imported by
# app.main (Tenant, ApiKey via auth middleware; Document via document_service;
# QueryLog via cost_tracker; EvalResult via admin router) would be registered,
# and any model not reachable from app.main would be missing from create_all.
import app.models  # noqa: F401


@pytest_asyncio.fixture(scope="session", autouse=True)
async def _create_test_schema():
    """Create all tables before the test session and drop them afterwards."""
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)


@pytest.fixture
def tenant_id():
    return uuid.uuid4()


@pytest.fixture
def sample_pdf_content():
    """Minimal valid PDF content for testing."""
    # Create a simple PDF-like content
    return b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << >> >>\nendobj\n4 0 obj\n<< /Length 44 >>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(Hello World) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000230 00000 n \ntrailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n326\n%%EOF"


@pytest.fixture
def sample_docx_content():
    """Create minimal DOCX bytes for testing."""
    from docx import Document as DocxDocument
    import io

    doc = DocxDocument()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph with some content for testing the extraction pipeline.")
    doc.add_heading("Second Section", level=2)
    doc.add_paragraph("This is the second section with more content.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_html_content():
    return b"""<!DOCTYPE html>
<html><head><title>Test Doc</title></head>
<body>
<h1>Main Title</h1>
<p>This is the first paragraph with important information.</p>
<h2>Second Section</h2>
<p>This is the second section with more details about the topic.</p>
</body></html>"""


@pytest.fixture
def sample_markdown_content():
    return b"""# Test Document

This is the introduction paragraph with some important content.

## Section One

This section contains details about the first topic. It has multiple sentences
to ensure proper chunking behavior.

## Section Two

This section covers the second topic with additional details.
"""

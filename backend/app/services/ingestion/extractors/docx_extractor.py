import io

from docx import Document as DocxDocument

from app.services.ingestion.extractors.base import BaseExtractor, ExtractedDocument, PageContent


class DocxExtractor(BaseExtractor):
    def extract(self, file_content: bytes, filename: str) -> ExtractedDocument:
        doc = DocxDocument(io.BytesIO(file_content))
        pages: list[PageContent] = []
        current_heading: str | None = None
        current_text_parts: list[str] = []
        section_index = 0

        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                # Save previous section
                if current_text_parts:
                    section_index += 1
                    pages.append(
                        PageContent(
                            text="\n".join(current_text_parts).strip(),
                            page_number=section_index,
                            section_heading=current_heading,
                        )
                    )
                    current_text_parts = []
                current_heading = para.text.strip()
            elif para.text.strip():
                current_text_parts.append(para.text.strip())

        # Last section
        if current_text_parts:
            section_index += 1
            pages.append(
                PageContent(
                    text="\n".join(current_text_parts).strip(),
                    page_number=section_index,
                    section_heading=current_heading,
                )
            )

        # If no headings were found, treat the whole doc as one page
        if not pages:
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if full_text:
                pages.append(PageContent(text=full_text, page_number=1))

        return ExtractedDocument(pages=pages, total_pages=len(pages), title=filename)
